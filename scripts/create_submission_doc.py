#!/usr/bin/env python3
"""Generate Word document for Innovation Track submission."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_submission_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('twodirect — Innovation Track Submission', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Track Selection
    doc.add_heading('Innovation Track ที่เลือก', level=1)
    doc.add_paragraph('✅ Track 2: Intelligent Retail Ecosystems', style='Intense Quote')
    doc.add_paragraph('เชื่อมต่อหน้าร้าน–หลังร้านด้วยระบบอัจฉริยะเพื่อเพิ่มประสิทธิภาพทางธุรกิจ')
    
    doc.add_paragraph()  # Spacer
    
    # Question 1
    doc.add_heading('1. ปัญหาหรือความท้าทายใดในอุตสาหกรรมการค้าปลีก ที่คุณต้องการแก้ไข หรือคุณมองเห็นโอกาสใหม่ใดที่ควรพัฒนา?', level=1)
    
    doc.add_heading('ปัญหาหลัก: ลูกค้าไม่สามารถรู้ได้ว่าสินค้าที่ต้องการมีอยู่ที่สาขาไหนใกล้บ้าน', level=2)
    
    doc.add_paragraph('ในปัจจุบัน ผู้บริโภคที่ต้องการซื้อสินค้าเฉพาะ เช่น สินค้าออกใหม่ สินค้า limited edition หรือสินค้าที่หายาก ต้องเผชิญกับปัญหาสำคัญดังนี้:')
    
    problems = [
        'เสียเวลาเดินทางไปหลายสาขาโดยไม่รู้ว่ามีสินค้าหรือไม่ — ลูกค้าต้องไป "สุ่ม" หาสินค้าตามร้านต่างๆ บางครั้งไป 3-4 สาขายังหาไม่เจอ ทำให้เสียเวลาและค่าเดินทาง',
        'แอปปัจจุบันไม่แสดงข้อมูลสต็อกรายสาขา — แอปของร้านค้าปลีกส่วนใหญ่บอกได้แค่ว่าร้านอยู่ที่ไหน แต่ไม่สามารถบอกได้ว่า "สินค้าชิ้นนี้" มีอยู่ที่สาขาไหนบ้าง',
        'ลูกค้ายอมแพ้และไม่ซื้อ — เมื่อหาไม่เจอหลายครั้ง ลูกค้าจะเลิกพยายามและไปซื้อสินค้าทดแทนหรือไม่ซื้อเลย ส่งผลให้ร้านค้าเสียยอดขาย',
        'ร้านค้าไม่สามารถสื่อสารกับลูกค้าที่กำลังหาสินค้า — แม้สินค้าจะมีอยู่ในสต็อก แต่ถ้าลูกค้าไม่รู้ ก็ไม่มีทางขายได้',
    ]
    for p in problems:
        doc.add_paragraph(p, style='List Bullet')
    
    doc.add_paragraph()
    opportunity = doc.add_paragraph()
    opportunity.add_run('โอกาสที่เห็น: ').bold = True
    opportunity.add_run('การเชื่อมต่อข้อมูลสต็อกสินค้ารายสาขากับระบบค้นหาอัจฉริยะ จะช่วยให้ลูกค้าหาสินค้าได้ตรงจุด และช่วยให้ร้านค้าเพิ่มยอดขายจากลูกค้าที่มี intent ชัดเจนอยู่แล้ว')
    
    # Question 2
    doc.add_heading('2. โปรดอธิบายแนวทางหรือวิธีการของคุณในการแก้ไขปัญหาหรือสนับสนุนโอกาสดังกล่าว เพื่อพัฒนาไอเดียของคุณให้สำเร็จ', level=1)
    
    doc.add_heading('twodirect — แพลตฟอร์มค้นหาสินค้าตามสาขาแบบ Real-time', level=2)
    
    doc.add_paragraph('แนวทางการแก้ปัญหาของเรามี 3 องค์ประกอบหลัก:')
    
    # Component 1
    comp1 = doc.add_paragraph()
    comp1.add_run('1. ระบบค้นหาสินค้าอัจฉริยะ (Intelligent Product Search)').bold = True
    doc.add_paragraph('ลูกค้าสามารถค้นหาสินค้าได้ 2 วิธี: พิมพ์ชื่อสินค้า หรือ ถ่ายรูป/อัพโหลดภาพสินค้า', style='List Bullet')
    doc.add_paragraph('ระบบใช้ AI Image Recognition (Google Cloud Vision) เพื่อระบุสินค้าจากภาพ', style='List Bullet')
    doc.add_paragraph('รองรับการค้นหาภาษาไทยและภาษาอังกฤษ', style='List Bullet')
    
    # Component 2
    comp2 = doc.add_paragraph()
    comp2.add_run('2. การเชื่อมต่อข้อมูลสต็อกรายสาขา (Branch-Level Inventory Integration)').bold = True
    doc.add_paragraph('เชื่อมต่อกับระบบ POS/Inventory ของพาร์ทเนอร์ค้าปลีกผ่าน API', style='List Bullet')
    doc.add_paragraph('แสดงข้อมูลสต็อกแบบ real-time หรือ near real-time ของแต่ละสาขา', style='List Bullet')
    doc.add_paragraph('ลูกค้าเห็นได้ทันทีว่าสาขาไหนใกล้บ้านมีสินค้าที่ต้องการ', style='List Bullet')
    
    # Component 3
    comp3 = doc.add_paragraph()
    comp3.add_run('3. ระบบนำทางและ Location-Based Services').bold = True
    doc.add_paragraph('แสดงระยะทางจากตำแหน่งลูกค้าไปยังแต่ละสาขาที่มีสินค้า', style='List Bullet')
    doc.add_paragraph('เรียงลำดับสาขาจากใกล้ไปไกล', style='List Bullet')
    doc.add_paragraph('กดปุ่มเดียวเพื่อนำทางไปยังสาขาที่เลือก', style='List Bullet')
    
    # Development phases
    doc.add_paragraph()
    phases = doc.add_paragraph()
    phases.add_run('ขั้นตอนการพัฒนา:').bold = True
    doc.add_paragraph('Phase 1: Validate ปัญหากับลูกค้าจริง และหาพาร์ทเนอร์ร้านค้าปลีกรายแรก', style='List Bullet')
    doc.add_paragraph('Phase 2: พัฒนา MVP และทดสอบกับ 50-100 สาขา', style='List Bullet')
    doc.add_paragraph('Phase 3: Launch สู่ตลาดและขยายจำนวนพาร์ทเนอร์', style='List Bullet')
    
    # Question 3
    doc.add_heading('3. ไอเดียของคุณแตกต่างจากที่มีอยู่ในตลาดหรือของคู่แข่งอย่างไร?', level=1)
    
    highlight = doc.add_paragraph()
    highlight.add_run('twodirect เป็นโซลูชันเดียวที่ตอบคำถาม: "สินค้านี้มีที่สาขาไหนใกล้ฉันบ้าง?"').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('การเปรียบเทียบกับคู่แข่ง:')
    
    competitors = [
        ('Grab / LINE MAN', 'ส่งสินค้ามาให้ที่บ้าน', 'twodirect ช่วยให้ลูกค้าไปหยิบเองได้ — เร็วกว่า ฟรีค่าส่ง'),
        ('Shopee / Lazada', 'สั่งออนไลน์ รอ 1-7 วัน', 'twodirect ได้ของวันนี้ ภายในชั่วโมงเดียว'),
        ('แอป 7-Eleven', 'บอกว่าร้านอยู่ที่ไหน แต่ไม่บอกว่าสินค้าอยู่สาขาไหน', 'twodirect บอกได้ว่าสินค้าชิ้นนั้นมีที่สาขาไหน'),
        ('โทรถามร้าน', 'เสียเวลา พนักงานอาจไม่รู้', 'twodirect ค้นหาได้ทันที ข้อมูลตรงจากระบบ'),
        ('ถามในโซเชียล', 'ช้า ไม่น่าเชื่อถือ', 'twodirect ข้อมูล real-time จากระบบหลังร้าน'),
    ]
    
    for comp, their_way, our_diff in competitors:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{comp}: ').bold = True
        p.add_run(f'{their_way} → {our_diff}')
    
    doc.add_paragraph()
    strengths = doc.add_paragraph()
    strengths.add_run('จุดแข็งที่แตกต่าง:').bold = True
    doc.add_paragraph('Branch-Level Visibility — เป็นเจ้าเดียวที่แสดงข้อมูลสต็อกระดับสาขา', style='List Bullet')
    doc.add_paragraph('Visual Search — ค้นหาด้วยภาพสำหรับลูกค้าที่ไม่รู้ชื่อสินค้า', style='List Bullet')
    doc.add_paragraph('Offline-First — เน้น pickup ที่ร้าน ไม่ใช่ delivery (เร็วกว่า+ฟรี)', style='List Bullet')
    doc.add_paragraph('Multi-Chain Potential — สามารถรวมหลายเชนในแอปเดียว', style='List Bullet')
    
    # Question 4
    doc.add_heading('4. Revenue Model ของคุณมีลักษณะอย่างไร? และมีการปรับเปลี่ยนหรือพัฒนาในรูปแบบใดเพื่อเพิ่มความเหมาะสมกับตลาดค้าปลีก?', level=1)
    
    doc.add_heading('Revenue Model แบบ 2 ด้าน: B2B + B2C', level=2)
    
    b2b = doc.add_paragraph()
    b2b.add_run('ด้าน B2B — รายได้จากพาร์ทเนอร์ร้านค้าปลีก:').bold = True
    doc.add_paragraph('Subscription Fee — ค่าสมาชิกรายเดือน/รายปี สำหรับการเชื่อมต่อ API และแสดงสินค้าบนแพลตฟอร์ม', style='List Bullet')
    doc.add_paragraph('Advertising Revenue — โฆษณาสินค้าในแอป แบบ promoted placement หรือ banner ads', style='List Bullet')
    doc.add_paragraph('Campaign Fee — ค่าโปรโมทสินค้าใหม่ หรือ flash sale ให้เข้าถึงลูกค้าเป้าหมาย', style='List Bullet')
    doc.add_paragraph('Performance Fee — Revenue share จากยอดขายที่เกิดจากการค้นหาผ่าน twodirect', style='List Bullet')
    
    doc.add_paragraph()
    b2c = doc.add_paragraph()
    b2c.add_run('ด้าน B2C — รายได้จากผู้ใช้งาน (Freemium Model):').bold = True
    doc.add_paragraph('Free — รัศมี 5 กิโลเมตร (ฟรี)', style='List Bullet')
    doc.add_paragraph('Plan 1 — รัศมี 15 กิโลเมตร (฿29/เดือน)', style='List Bullet')
    doc.add_paragraph('Plan 2 — ทั่วประเทศ (฿59/เดือน)', style='List Bullet')
    
    doc.add_paragraph()
    adapt = doc.add_paragraph()
    adapt.add_run('การปรับให้เหมาะกับตลาดค้าปลีก:').bold = True
    doc.add_paragraph('Value-Based Pricing — พาร์ทเนอร์จ่ายตามมูลค่าที่ได้รับ (foot traffic, sales)', style='List Bullet')
    doc.add_paragraph('Low Entry Barrier — ลูกค้าใช้ฟรีได้ก่อน สร้าง adoption ก่อนขาย premium', style='List Bullet')
    doc.add_paragraph('Win-Win Model — ร้านค้าได้ยอดขาย ลูกค้าประหยัดเวลา twodirect ได้รายได้', style='List Bullet')
    
    # Question 5
    doc.add_heading('5. ไอเดียหรือโซลูชันของคุณมีศักยภาพในการขยายไปสู่ตลาดระดับสากล (Global) อย่างไร?', level=1)
    
    global1 = doc.add_paragraph()
    global1.add_run('ศักยภาพการขยายสู่ตลาดสากล:').bold = True
    
    doc.add_paragraph()
    uni = doc.add_paragraph()
    uni.add_run('1. ปัญหานี้เป็น Universal Problem').bold = True
    doc.add_paragraph('ทุกประเทศมีร้านค้าปลีกที่มีหลายสาขา', style='List Bullet')
    doc.add_paragraph('ลูกค้าทุกที่ต้องการหาสินค้าได้ตรงจุด', style='List Bullet')
    doc.add_paragraph('ตลาดหลักที่มีศักยภาพ: เอเชียตะวันออกเฉียงใต้, ญี่ปุ่น, เกาหลี, อินเดีย', style='List Bullet')
    
    doc.add_paragraph()
    plan = doc.add_paragraph()
    plan.add_run('2. แผนการขยายระยะยาว:').bold = True
    doc.add_paragraph('Phase 1 (ปี 2026): ครองตลาดไทย — 7-Eleven, Lawson, FamilyMart', style='List Bullet')
    doc.add_paragraph('Phase 2 (ปี 2027): ขยายสู่ ASEAN — เวียดนาม, มาเลเซีย, อินโดนีเซีย', style='List Bullet')
    doc.add_paragraph('Phase 3 (ปี 2028+): ขยายสู่ตลาดเอเชีย — ญี่ปุ่น, เกาหลี, อินเดีย', style='List Bullet')
    
    doc.add_paragraph()
    risk = doc.add_paragraph()
    risk.add_run('3. กลยุทธ์ลดความเสี่ยงในการขยายธุรกิจ:').bold = True
    doc.add_paragraph('Proven Model First — พิสูจน์โมเดลในไทยก่อน สร้าง case study ที่แข็งแกร่ง', style='List Bullet')
    doc.add_paragraph('Same Chain Expansion — ขยายตาม chain ที่มีอยู่หลายประเทศ (เช่น 7-Eleven มีใน 19 ประเทศ)', style='List Bullet')
    doc.add_paragraph('Local Partnership — หาพาร์ทเนอร์ท้องถิ่นในแต่ละประเทศ ไม่ลงทุนเองทั้งหมด', style='List Bullet')
    doc.add_paragraph('Tech-First Approach — Platform เดียวใช้ได้หลายประเทศ เปลี่ยนแค่ภาษาและ API connection', style='List Bullet')
    doc.add_paragraph('Franchise Model — License โมเดลให้ผู้ประกอบการท้องถิ่น ลดความเสี่ยงด้านการลงทุน', style='List Bullet')
    
    doc.add_paragraph()
    scale = doc.add_paragraph()
    scale.add_run('4. Scalability ของเทคโนโลยี:').bold = True
    doc.add_paragraph('Cloud-based infrastructure (AWS/GCP) ขยายได้ทั่วโลก', style='List Bullet')
    doc.add_paragraph('API-first architecture — เชื่อมต่อกับระบบค้าปลีกได้ทุกรูปแบบ', style='List Bullet')
    doc.add_paragraph('AI/ML ที่ปรับได้ตาม product catalog ของแต่ละประเทศ', style='List Bullet')
    
    # Save document
    doc.save('note/twodirect-innovation-track-submission.docx')
    print('✅ Document saved: note/twodirect-innovation-track-submission.docx')

if __name__ == '__main__':
    create_submission_document()


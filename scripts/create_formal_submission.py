#!/usr/bin/env python3
"""Generate formal Word document for Innovation Track submission."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_formal_submission():
    doc = Document()
    
    # Title
    title = doc.add_heading('twodirect — Innovation Track Submission', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Question 1
    doc.add_heading('1. ปัญหาหรือความท้าทายใดในอุตสาหกรรมการค้าปลีก ที่คุณต้องการแก้ไข หรือคุณมองเห็นโอกาสใหม่ใดที่ควรพัฒนา?', level=1)
    
    doc.add_heading('ปัญหาหลัก: ลูกค้าไม่สามารถรู้ได้ว่าสินค้าที่ต้องการมีอยู่ที่สาขาไหน', level=2)
    
    doc.add_paragraph('ในปัจจุบัน ผู้บริโภคที่ต้องการซื้อสินค้าเฉพาะ เช่น สินค้าออกใหม่ สินค้า Limited Edition หรือสินค้าที่หายาก ต้องเผชิญกับปัญหาสำคัญดังนี้:')
    
    doc.add_paragraph('เสียเวลาเดินทางไปหลายสาขาโดยไม่รู้ว่ามีสินค้าหรือไม่ — ลูกค้าต้องไป "สุ่ม" หาสินค้าตามร้านต่างๆ บางครั้งไป 3-4 สาขายังหาไม่เจอ ทำให้เสียเวลาและค่าเดินทาง', style='List Bullet')
    
    # FORMALIZED red text
    doc.add_paragraph('แอปพลิเคชันปัจจุบันไม่แสดงข้อมูลสต็อกรายสาขา — แอปพลิเคชันของร้านค้าปลีกส่วนใหญ่สามารถบอกได้เพียงว่าร้านตั้งอยู่ที่ใด แต่ไม่สามารถระบุได้ว่า "สินค้าชิ้นนี้" มีจำหน่ายที่สาขาใดบ้าง นอกจากนี้ บางแอปพลิเคชันแสดงเฉพาะสาขาใกล้เคียงที่สินค้าหมด ทำให้ลูกค้าต้องติดต่อสอบถามผู้ดูแลระบบว่าสินค้ามีจำหน่ายที่สาขาใด ส่งผลให้เพิ่มภาระงานให้กับเจ้าหน้าที่ในการตอบคำถามเหล่านี้ในแต่ละวัน', style='List Bullet')
    
    doc.add_paragraph('ลูกค้ายอมแพ้และไม่ซื้อ — เมื่อหาไม่เจอหลายครั้ง ลูกค้าจะเลิกพยายามและไปซื้อสินค้าทดแทนหรือไม่ซื้อเลย ส่งผลให้ร้านค้าเสียยอดขาย', style='List Bullet')
    
    doc.add_paragraph('ร้านค้าไม่สามารถสื่อสารกับลูกค้าที่กำลังหาสินค้า — แม้สินค้าจะมีอยู่ในสต็อก แต่ถ้าลูกค้าไม่รู้ ก็ไม่มีทางขายได้', style='List Bullet')
    
    doc.add_paragraph()
    opp = doc.add_paragraph()
    opp.add_run('โอกาสที่เห็น: ').bold = True
    opp.add_run('การเชื่อมต่อข้อมูลสต็อกสินค้ารายสาขากับระบบค้นหาอัจฉริยะ จะช่วยให้ลูกค้าหาสินค้าได้ตรงจุด และช่วยให้ร้านค้าเพิ่มยอดขายจากลูกค้าที่มีความต้องการซื้อที่ชัดเจนอยู่แล้ว')
    
    # Question 2
    doc.add_heading('2. โปรดอธิบายแนวทางหรือวิธีการของคุณในการแก้ไขปัญหาหรือสนับสนุนโอกาสดังกล่าว เพื่อพัฒนาไอเดียของคุณให้สำเร็จ', level=1)
    
    doc.add_heading('twodirect — แพลตฟอร์มค้นหาสินค้าตามสาขาแบบ Real-time', level=2)
    
    # FORMALIZED red text - concept explanation
    concept = doc.add_paragraph()
    concept.add_run('แนวคิดหลัก: ').bold = True
    concept.add_run('ในยุคปัจจุบัน ผู้บริโภคคุ้นเคยกับการใช้เทคโนโลยีเพื่อตอบโจทย์ความต้องการต่างๆ ไม่ว่าจะเป็นการค้นหาข้อมูลผ่าน Google การนำทางผ่าน Google Maps การสั่งซื้อสินค้าออนไลน์ผ่าน Shopee, TikTok Shop หรือ Lazada เพื่อรับส่วนลด หรือการสั่งอาหารผ่าน Grab และ LINE MAN ทีมงานเล็งเห็นความท้าทายในกระบวนการค้นหาสินค้าที่ร้านค้าปลีก จึงพัฒนาโซลูชันที่ช่วยลดเวลาในการค้นหา เพิ่มเวลาให้กับผู้บริโภค และทำให้ได้รับสินค้าที่ต้องการอย่างแม่นยำ เปรียบเสมือนการค้นหาโทรศัพท์มือถือที่ลืมไว้ในบ้านแต่ไม่ทราบตำแหน่งที่แน่นอน ซึ่งสามารถใช้ฟังก์ชัน Find My Phone เพื่อระบุพิกัดได้ทันที')
    
    doc.add_paragraph()
    doc.add_paragraph('แนวทางการแก้ปัญหาของเรามี 3 องค์ประกอบหลัก:')
    
    # Component 1
    c1 = doc.add_paragraph()
    c1.add_run('1. ระบบค้นหาสินค้าอัจฉริยะ (Intelligent Product Search)').bold = True
    doc.add_paragraph('ลูกค้าสามารถค้นหาสินค้าได้ 2 วิธี: พิมพ์ชื่อสินค้า หรือ ถ่ายรูป/อัพโหลดภาพสินค้า', style='List Bullet')
    doc.add_paragraph('ระบบใช้ AI Image Recognition (Google Cloud Vision) เพื่อระบุสินค้าจากภาพ', style='List Bullet')
    # FORMALIZED
    doc.add_paragraph('รองรับการค้นหาภาษาไทยและภาษาอังกฤษ โดยในอนาคตจะขยายการรองรับหลายภาษาเพิ่มเติม เนื่องจากเครือข่ายร้านค้าปลีกอย่าง 7-Eleven ดำเนินกิจการในหลายประเทศทั่วโลก ทำให้สามารถให้บริการลูกค้าได้หลากหลายกลุ่ม โดยใช้ต้นแบบการให้บริการหลายภาษาจาก Netflix', style='List Bullet')
    
    # Component 2
    c2 = doc.add_paragraph()
    c2.add_run('2. การเชื่อมต่อข้อมูลสต็อกรายสาขา (Branch-Level Inventory Integration)').bold = True
    doc.add_paragraph('เชื่อมต่อกับระบบ POS/Inventory ของพาร์ทเนอร์ค้าปลีกผ่าน API', style='List Bullet')
    doc.add_paragraph('แสดงข้อมูลสต็อกแบบ Real-time หรือ Near Real-time ของแต่ละสาขา', style='List Bullet')
    doc.add_paragraph('ลูกค้าเห็นได้ทันทีว่าสาขาไหนใกล้บ้านมีสินค้าที่ต้องการ', style='List Bullet')
    
    # Component 3
    c3 = doc.add_paragraph()
    c3.add_run('3. ระบบนำทางและ Location-Based Services').bold = True
    doc.add_paragraph('แสดงระยะทางจากตำแหน่งลูกค้าไปยังแต่ละสาขาที่มีสินค้า', style='List Bullet')
    doc.add_paragraph('เรียงลำดับสาขาจากใกล้ไปไกล', style='List Bullet')
    doc.add_paragraph('กดปุ่มเดียวเพื่อนำทางไปยังสาขาที่เลือก', style='List Bullet')
    # FORMALIZED
    doc.add_paragraph('สามารถนำทางไปยังแอปพลิเคชันหลักของพาร์ทเนอร์เพื่อดำเนินการสั่งซื้อได้โดยตรง', style='List Bullet')
    
    doc.add_paragraph()
    dev = doc.add_paragraph()
    dev.add_run('ขั้นตอนการพัฒนา:').bold = True
    doc.add_paragraph('Phase 1: ตรวจสอบปัญหากับกลุ่มลูกค้าเป้าหมาย และหาพาร์ทเนอร์ร้านค้าปลีกรายแรก', style='List Bullet')
    doc.add_paragraph('Phase 2: พัฒนา MVP และทดสอบกับ 50-100 สาขา', style='List Bullet')
    # FORMALIZED
    doc.add_paragraph('Phase 3: เปิดตัวสู่ตลาดและขยายจำนวนพาร์ทเนอร์ภายในประเทศ', style='List Bullet')
    doc.add_paragraph('Phase 4: ขยายการดำเนินงานสู่ตลาดต่างประเทศและเพิ่มจำนวนพาร์ทเนอร์ในระดับสากล', style='List Bullet')
    
    # Question 3
    doc.add_heading('3. ไอเดียของคุณแตกต่างจากที่มีอยู่ในตลาดหรือของคู่แข่งอย่างไร?', level=1)
    
    # FORMALIZED
    highlight = doc.add_paragraph()
    highlight.add_run('twodirect เป็นโซลูชันเดียวที่ตอบคำถาม: "สินค้านี้มีที่สาขาไหนใกล้ฉันบ้าง?" ').bold = True
    highlight.add_run('พร้อมทั้งมีระบบส่วนลดพิเศษที่ช่วยดึงดูดให้ลูกค้ากลับมาใช้งานและเกิดการซื้อซ้ำ')
    
    doc.add_paragraph()
    doc.add_paragraph('การเปรียบเทียบกับคู่แข่ง:')
    
    doc.add_paragraph('Grab / LINE MAN: ให้บริการจัดส่งสินค้าถึงที่พักอาศัย → twodirect ช่วยให้ลูกค้าสามารถไปรับสินค้าด้วยตนเองได้ ซึ่งรวดเร็วกว่าและไม่มีค่าจัดส่ง', style='List Bullet')
    # FORMALIZED
    doc.add_paragraph('Shopee / Lazada / TikTok Shop: สั่งซื้อออนไลน์และรอรับสินค้า 1-7 วัน → twodirect ทำให้ลูกค้าได้รับสินค้าภายในวันเดียวกัน ตามระยะเวลาที่ใช้ในการเดินทางไปยังสาขาที่ต้องการ', style='List Bullet')
    # FORMALIZED  
    doc.add_paragraph('แอปพลิเคชัน 7-Eleven: แสดงเฉพาะตำแหน่งที่ตั้งของร้าน แต่ไม่ระบุว่าสินค้าใดมีจำหน่ายที่สาขาใด → twodirect สามารถระบุได้ว่าสินค้าชิ้นนั้นมีจำหน่ายที่สาขาใดบ้าง พร้อมทั้งมีส่วนลดพิเศษในแอปพลิเคชันนอกเหนือจากโปรโมชันของสาขา', style='List Bullet')
    doc.add_paragraph('โทรสอบถามร้าน: เสียเวลาและพนักงานอาจไม่ทราบข้อมูล → twodirect ค้นหาได้ทันที โดยดึงข้อมูลโดยตรงจากระบบ', style='List Bullet')
    doc.add_paragraph('สอบถามในโซเชียลมีเดีย: ใช้เวลานานและข้อมูลไม่น่าเชื่อถือ → twodirect ให้ข้อมูลแบบ Real-time จากระบบหลังร้าน', style='List Bullet')
    
    doc.add_paragraph()
    str_title = doc.add_paragraph()
    str_title.add_run('จุดแข็งที่แตกต่าง:').bold = True
    doc.add_paragraph('Branch-Level Visibility — เป็นรายเดียวที่แสดงข้อมูลสต็อกระดับสาขา', style='List Bullet')
    doc.add_paragraph('Visual Search — ค้นหาด้วยภาพสำหรับลูกค้าที่ไม่ทราบชื่อสินค้า', style='List Bullet')
    # FORMALIZED
    doc.add_paragraph('Offline-First — เน้นการไปรับสินค้าที่ร้านด้วยตนเอง ไม่ใช่การจัดส่ง (รวดเร็วกว่าและไม่มีค่าใช้จ่าย) ทำให้ลูกค้าไม่เสียเที่ยว เพราะเวลามีคุณค่าและไม่สามารถซื้อหาได้', style='List Bullet')
    # FORMALIZED
    doc.add_paragraph('Multi-Chain Potential — สามารถรวมหลายเครือข่ายร้านค้าในแอปพลิเคชันเดียว ให้บริการแบบ One-Stop Service', style='List Bullet')

    doc.add_paragraph()
    # FORMALIZED - additional differentiation
    add_diff = doc.add_paragraph()
    add_diff.add_run('ความยืดหยุ่นในการเป็นพาร์ทเนอร์: ').bold = True
    add_diff.add_run('twodirect ไม่จำกัดเฉพาะประเภทธุรกิจใดธุรกิจหนึ่ง แต่เปิดรับทุกธุรกิจที่มีหลายสาขาเป็นพาร์ทเนอร์ ไม่ว่าจะเป็นร้านอาหาร ร้านเสื้อผ้า หรือร้านของใช้ทั่วไป')
    
    doc.add_paragraph()
    # FORMALIZED - discount strategy
    discount = doc.add_paragraph()
    discount.add_run('กลยุทธ์ส่วนลดเพื่อดึงดูดการใช้งาน: ').bold = True
    discount.add_run('twodirect มอบส่วนลดพิเศษนอกเหนือจากโปรโมชันในแอปพลิเคชันหลักของแต่ละกิจการ อย่างไรก็ตาม หากลูกค้าเลือกซื้อแบบ Offline จะไม่ได้รับส่วนลดดังกล่าว แต่ระบบยังคงให้ข้อมูลตำแหน่งสินค้าเพื่อเพิ่มโอกาสในการเข้าถึงร้านค้า กลยุทธ์นี้ส่งเสริมให้ลูกค้าหันมาสั่งซื้อผ่านแอปพลิเคชัน ดึงดูดการใช้งานซ้ำ ให้ความคุ้มค่าทั้งด้านราคาและเวลา ทำให้ลูกค้าสามารถนำเวลาที่ประหยัดได้ไปทำกิจกรรมอื่นๆ โดยยังได้รับสินค้าที่ต้องการ')
    
    # Question 4
    doc.add_heading('4. Revenue Model ของคุณมีลักษณะอย่างไร? และมีการปรับเปลี่ยนหรือพัฒนาในรูปแบบใดเพื่อเพิ่มความเหมาะสมกับตลาดค้าปลีก?', level=1)
    
    doc.add_heading('Revenue Model แบบ 2 ด้าน: B2B + B2C', level=2)
    
    b2b = doc.add_paragraph()
    b2b.add_run('ด้าน B2B — รายได้จากพาร์ทเนอร์ร้านค้าปลีก:').bold = True
    doc.add_paragraph('Subscription Fee — ค่าสมาชิกรายเดือน/รายปี สำหรับการเชื่อมต่อ API และแสดงสินค้าบนแพลตฟอร์ม', style='List Bullet')
    doc.add_paragraph('Advertising Revenue — โฆษณาสินค้าในแอปพลิเคชัน แบบ Promoted Placement หรือ Banner Ads', style='List Bullet')
    doc.add_paragraph('Campaign Fee — ค่าโปรโมทสินค้าใหม่ หรือ Flash Sale ให้เข้าถึงลูกค้าเป้าหมาย', style='List Bullet')
    doc.add_paragraph('Performance Fee — Revenue Share จากยอดขายที่เกิดจากการค้นหาผ่าน twodirect', style='List Bullet')

    doc.add_paragraph()
    b2c = doc.add_paragraph()
    b2c.add_run('ด้าน B2C — รายได้จากผู้ใช้งาน (Freemium Model):').bold = True
    doc.add_paragraph('Free — รัศมีการค้นหา 5 กิโลเมตร (ไม่มีค่าใช้จ่าย)', style='List Bullet')
    doc.add_paragraph('Plan 1 — รัศมีการค้นหา 15 กิโลเมตร (฿29/เดือน)', style='List Bullet')
    doc.add_paragraph('Plan 2 — ค้นหาได้ทั่วประเทศ (฿59/เดือน)', style='List Bullet')

    doc.add_paragraph()
    # FORMALIZED - pricing note
    pricing_note = doc.add_paragraph()
    pricing_note.add_run('หมายเหตุด้านราคา: ').bold = True
    pricing_note.add_run('โครงสร้างราคาอาจมีการปรับเปลี่ยนตามความเหมาะสม เนื่องจากระบบมีการให้ส่วนลดแก่ผู้ใช้งาน ดังนั้นจำเป็นต้องบริหารจัดการต้นทุนให้อยู่ในระดับต่ำเพื่อให้สามารถทำกำไรได้ในราคาที่กำหนด')
    
    doc.add_paragraph()
    adapt = doc.add_paragraph()
    adapt.add_run('การปรับให้เหมาะกับตลาดค้าปลีก:').bold = True
    doc.add_paragraph('Value-Based Pricing — พาร์ทเนอร์จ่ายตามมูลค่าที่ได้รับ (Foot Traffic และยอดขาย)', style='List Bullet')
    doc.add_paragraph('Low Entry Barrier — ลูกค้าใช้งานฟรีได้ก่อน เพื่อสร้างฐานผู้ใช้ก่อนเสนอขายแพ็กเกจ Premium', style='List Bullet')
    doc.add_paragraph('Win-Win Model — ร้านค้าได้ยอดขายเพิ่ม ลูกค้าประหยัดเวลา และ twodirect ได้รายได้', style='List Bullet')
    
    # Question 5
    doc.add_heading('5. ไอเดียหรือโซลูชันของคุณมีศักยภาพในการขยายไปสู่ตลาดระดับสากล (Global) อย่างไร?', level=1)
    
    global_title = doc.add_paragraph()
    global_title.add_run('ศักยภาพการขยายสู่ตลาดสากล:').bold = True
    
    doc.add_paragraph()
    uni = doc.add_paragraph()
    uni.add_run('1. ปัญหานี้เป็น Universal Problem').bold = True
    doc.add_paragraph('ทุกประเทศมีร้านค้าปลีกที่มีหลายสาขา', style='List Bullet')
    doc.add_paragraph('ลูกค้าทุกที่ต้องการหาสินค้าได้ตรงจุด', style='List Bullet')
    doc.add_paragraph('ตลาดหลักที่มีศักยภาพ: เอเชียตะวันออกเฉียงใต้, ญี่ปุ่น, เกาหลี, อินเดีย', style='List Bullet')
    
    doc.add_paragraph()
    plan = doc.add_paragraph()
    plan.add_run('2. แผนการขยายระยะยาว:').bold = True
    # FORMALIZED
    doc.add_paragraph('Phase 1 (ปี 2026): ครองตลาดประเทศไทย — 7-Eleven, Makro, Lotus\'s และ Tops', style='List Bullet')
    doc.add_paragraph('Phase 2 (ปี 2028): ขยายสู่ภูมิภาคอาเซียน — เวียดนาม, มาเลเซีย, อินโดนีเซีย', style='List Bullet')
    doc.add_paragraph('Phase 3 (ปี 2029 เป็นต้นไป): ขยายสู่ตลาดเอเชีย — ญี่ปุ่น, เกาหลี, อินเดีย', style='List Bullet')
    
    doc.add_paragraph()
    risk = doc.add_paragraph()
    risk.add_run('3. กลยุทธ์ลดความเสี่ยงในการขยายธุรกิจ:').bold = True
    doc.add_paragraph('Proven Model First — พิสูจน์โมเดลธุรกิจในประเทศไทยก่อน เพื่อสร้างกรณีศึกษาที่แข็งแกร่ง', style='List Bullet')
    # FORMALIZED
    doc.add_paragraph('Same Chain Expansion — ขยายธุรกิจตามเครือข่ายร้านค้าที่ดำเนินกิจการในหลายประเทศ (เช่น 7-Eleven ซึ่งมีสาขาในหลากหลายประเทศทั่วโลก)', style='List Bullet')
    doc.add_paragraph('Local Partnership — หาพาร์ทเนอร์ท้องถิ่นในแต่ละประเทศ เพื่อลดการลงทุนด้วยตนเองทั้งหมด', style='List Bullet')
    doc.add_paragraph('Tech-First Approach — แพลตฟอร์มเดียวใช้งานได้หลายประเทศ เปลี่ยนเฉพาะภาษาและการเชื่อมต่อ API', style='List Bullet')
    doc.add_paragraph('Franchise Model — ให้สิทธิ์ใช้งานโมเดลธุรกิจแก่ผู้ประกอบการท้องถิ่น เพื่อลดความเสี่ยงด้านการลงทุน', style='List Bullet')
    
    doc.add_paragraph()
    scale = doc.add_paragraph()
    scale.add_run('4. Scalability ของเทคโนโลยี:').bold = True
    doc.add_paragraph('Cloud-Based Infrastructure (AWS/GCP) สามารถขยายการให้บริการได้ทั่วโลก', style='List Bullet')
    doc.add_paragraph('API-First Architecture — เชื่อมต่อกับระบบค้าปลีกได้ทุกรูปแบบ', style='List Bullet')
    doc.add_paragraph('AI/ML ที่ปรับแต่งได้ตาม Product Catalog ของแต่ละประเทศ', style='List Bullet')
    
    # Save
    doc.save('note/twodirect-innovation-track-submission-formal.docx')
    print('✅ Formal document saved: note/twodirect-innovation-track-submission-formal.docx')

if __name__ == '__main__':
    create_formal_submission()

